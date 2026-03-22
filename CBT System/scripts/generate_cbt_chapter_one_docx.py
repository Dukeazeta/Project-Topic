from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TOPIC_DIR = Path(__file__).resolve().parents[1]
INPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 01 Draft.md"
OUTPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 01 - Azeta Duke.docx"


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_list_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_term_paragraph(doc: Document, line: str) -> None:
    term, meaning = line.split(":", 1)
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(term + ":")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run2 = paragraph.add_run(meaning)
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2.font.size = Pt(12)


def main() -> None:
    doc = Document()
    set_default_font(doc)
    set_margins(doc)

    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
        elif line[:2].isdigit() and line[1] == ".":
            add_list_paragraph(doc, line)
        elif line.startswith("**") and ":**" in line:
            cleaned = line.replace("**", "")
            add_term_paragraph(doc, cleaned)
        else:
            add_body_paragraph(doc, line)

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
