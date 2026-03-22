from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TOPIC_DIR = Path(__file__).resolve().parents[1]
ROOT_DIR = TOPIC_DIR.parent
DOC_PATH = ROOT_DIR / "Shared" / "References" / "Project Proposal.docx"
BACKUP_PATH = ROOT_DIR / "Shared" / "References" / "Project Proposal.original.docx"
OUTPUT_PATH = TOPIC_DIR / "Proposal - Azeta Duke.docx"


def has_drawing(run) -> bool:
    xml = run._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def set_para_text(paragraph, text: str) -> None:
    text_runs = [run for run in paragraph.runs if not has_drawing(run)]
    if text_runs:
        text_runs[0].text = text
        for run in text_runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_para(paragraph) -> None:
    set_para_text(paragraph, "")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def main() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)

    # Title page
    set_para_text(doc.paragraphs[5], "Name: AZETA DUKE")
    set_para_text(
        doc.paragraphs[6],
        "Matric Number: COS/8650/2021 Supervisor/Mentor: Dr NWOZOR BLESSING Department: COMPUTER SCIENCE",
    )
    set_para_text(
        doc.paragraphs[7],
        "Institution: FEDERAL UNIVERSITY OF PETROLEUM RESOURCES Academic Session: 2025/2026",
    )
    set_para_text(
        doc.paragraphs[10],
        "Project Title: Design and Implementation of a Mobile-First Anti-Cheat CBT Platform for the Department of Computer Science, Federal University of Petroleum Resources, Effurun",
    )

    # Executive summary
    set_para_text(
        doc.paragraphs[14],
        "Computer-Based Testing (CBT) has become a common way of conducting examinations in many schools because it makes exams easier to organize, faster to mark, and simpler to manage. Even with these benefits, many CBT systems still have important weaknesses. Some are not convenient to use on mobile phones, while others do not have enough protection against common examination malpractice such as tab switching, copy and paste attempts, right-clicking, and poor session monitoring. Because of these weaknesses, many students and lecturers may still have doubts about the fairness of digital examinations.",
    )
    set_para_text(
        doc.paragraphs[16],
        "This project focuses on the design and implementation of a mobile-first anti-cheat CBT platform for the Department of Computer Science, Federal University of Petroleum Resources, Effurun. The system will be built as a standalone web platform for CBT activities. It will support exam creation, question import, student management, session monitoring, result export, secure student login, timed test taking, and practical anti-cheat controls such as fullscreen enforcement, tab switch detection, copy and paste blocking, right-click blocking, developer tools shortcut detection, warning logs, and automatic session termination after repeated violations. The expected result is a mobile-friendly CBT platform that improves exam integrity, supports easier administration, and gives staff and students a better digital examination experience.",
    )

    # Background study
    background_texts = [
        "Digital technology has changed the way teaching, learning, and assessment are done in higher institutions. One clear example of this is the use of Computer-Based Testing. CBT has become more popular because it reduces the stress of manual marking, improves record keeping, and helps schools handle examinations for many students at the same time. In Nigeria, the use of CBT has moved from large external examinations into universities and other tertiary institutions, although the success of implementation still differs from one institution to another (Adamu, 2024).",
        "Students are gradually becoming more comfortable with CBT because it can be faster, cleaner, and easier to manage than paper-based examinations. Tella and Bashorun (2012) reported that students in the Nigerian university environment can show a positive attitude toward CBT when the system is reliable and fair. Okocha (2022) also showed that student perception has a strong effect on whether CBT is accepted or not. This means that a good CBT platform should not only work well technically, but should also be simple to use and easy to trust.",
        "Another important issue is access to devices. In many Nigerian higher institutions, students depend heavily on smartphones to access online resources. Because of this, a modern CBT system should not be designed only for laptops and desktop computers. A mobile-first approach is important because it makes sure the system works well on smaller screens from the beginning instead of treating phone users as secondary users. Studies on mobile learning in higher education show that students are more willing to use mobile educational systems when the system is easy to use and matches their learning habits (Cheon et al., 2012; Naveed et al., 2023).",
        "Even with the advantages of CBT, exam integrity is still a major concern. Online assessments can be weakened by poor monitoring, weak session control, and limited anti-cheat enforcement. Students may try to switch tabs, copy questions, use outside help, or take advantage of weak invigilation. Holden et al. (2021) explained that academic integrity is one of the biggest concerns in online assessment environments. This shows that a useful CBT platform must combine usability with realistic integrity controls.",
        "In the Department of Computer Science, Federal University of Petroleum Resources, Effurun, there is a need for a CBT system that is department-focused, easy to access on mobile devices, and equipped with practical anti-cheat measures. This project addresses that need by proposing a standalone CBT platform tailored to the department's examination workflow.",
    ]
    for index, text in zip([19, 21, 23, 25, 27], background_texts):
        set_para_text(doc.paragraphs[index], text)
    clear_para(doc.paragraphs[29])

    # Problem statement
    problem_texts = [
        "Although CBT systems are now widely used in educational environments, many of them still do not properly combine mobile accessibility with practical anti-cheat controls. Some systems are hard to use on phones, while others do not provide enough monitoring for student sessions during examinations. Because of this, digital examinations may still be exposed to malpractice through tab switching, leaving fullscreen mode, copy and paste behavior, right-clicking, weak answer tracking, and poor administrative control over live exam sessions.",
        "In a department-level setting such as the Department of Computer Science, FUPRE, this creates a serious challenge. A system may allow students to take examinations digitally, but if it cannot properly manage exam time windows, session behavior, warnings, and submission control, the fairness and credibility of the assessment process become questionable.",
        "Also, since many students depend on mobile devices, a CBT platform that is not mobile-friendly may reduce usability and create access problems. The problem this project addresses therefore is the absence of a focused department-level CBT platform that is both mobile-first and supported by practical browser-based anti-cheat features for better examination integrity and administrative control.",
    ]
    for index, text in zip([32, 34, 35], problem_texts):
        set_para_text(doc.paragraphs[index], text)

    # Aim / objective
    set_para_text(
        doc.paragraphs[38],
        "To design and implement a mobile-first anti-cheat CBT platform for the Department of Computer Science, Federal University of Petroleum Resources, Effurun. The specific objectives of the study are as follows:",
    )
    objective_texts = [
        "To build a standalone CBT platform with separate interfaces for administrators and students.",
        "To provide support for exam creation, question management, and student management.",
        "To implement secure exam access with student login and session tracking.",
        "To implement browser-based anti-cheat controls such as fullscreen monitoring, tab switch detection, copy and paste blocking, right-click blocking, and warning logging.",
        "To support timed examinations, automatic submission, and violation-based session termination.",
        "To provide result handling, session monitoring, and administrative control tools such as pause, resume, and final submit.",
    ]
    objective_paragraphs = [doc.paragraphs[39], doc.paragraphs[40], doc.paragraphs[41]]
    for paragraph, text in zip(objective_paragraphs, objective_texts[:3]):
        set_para_text(paragraph, text)

    # Conceptual review
    set_para_text(doc.paragraphs[43], "CONCEPTUAL REVIEW")
    conceptual_texts = [
        "Computer-Based Testing is an examination method in which questions are delivered and answered through a digital system instead of paper. CBT improves speed, record keeping, and automatic grading for objective questions. It also helps institutions manage exam data more efficiently. However, for CBT to truly support fair academic assessment, the platform used must be dependable, secure, and easy to use.",
        "A mobile-first system is designed with phone users in mind from the beginning. This means the layout, navigation, controls, and content are first planned for smaller screens before being expanded for bigger screens. In this project, mobile-first design is important because it supports better access for students who rely on phones as their main digital device (Han & Shin, 2016; Naveed et al., 2023).",
        "Academic integrity means fairness, honesty, and trust in the assessment process. In online examinations, this means making sure that students take tests under acceptable conditions and that their results reflect their real performance. Holden et al. (2021) noted that online testing creates new integrity concerns because students are no longer always under the same physical supervision found in traditional examination halls.",
        "Session control refers to the ability of the system to track and manage each student's exam attempt from start to finish. This includes recording when the session starts, how long it lasts, whether it is paused, whether it has been submitted, and whether suspicious actions occur during the attempt. Good session control helps administrators monitor ongoing exams and respond quickly when problems come up.",
        "Browser-based anti-cheat monitoring uses the web browser to detect suspicious actions during an exam. Examples include detecting when a student leaves the exam tab, exits fullscreen mode, tries to copy or paste content, right-clicks, or attempts to open developer tools. These actions do not make cheating impossible, but they can discourage it and improve monitoring. In this project, anti-cheat is treated as practical prevention and detection, not as a claim that cheating can be removed completely.",
    ]
    for index, text in zip([44, 46, 48, 50, 52], conceptual_texts):
        set_para_text(doc.paragraphs[index], text)
    clear_para(doc.paragraphs[54])

    # Literature review tables
    literature_rows = [
        [
            "Adamu (2024), Computer-Based Testing Implementation in Nigeria: Successes and Challenges from Historical Perspectives",
            "To examine the growth and challenges of CBT implementation in Nigeria.",
            "Historical and analytical review of CBT implementation trends.",
            "Provides useful Nigerian context and identifies practical implementation challenges.",
            "Does not focus on mobile-first delivery or detailed anti-cheat controls.",
        ],
        [
            "Okocha (2022), Student Perception of Computer-Based Testing in Kwara State, Nigeria",
            "To study how students perceive CBT systems.",
            "Student perception study.",
            "Useful for understanding student trust, usability, and acceptance of CBT.",
            "Limited emphasis on technical design and session security.",
        ],
        [
            "Holden, Norris, and Kuhlmeier (2021), Academic Integrity in Online Assessment: A Research Review",
            "To review academic integrity issues in online assessment.",
            "Research review.",
            "Strong foundation for explaining integrity concerns in online exams.",
            "Broad review; not centered on a department-level CBT solution.",
        ],
        [
            "Dendir and Maxwell (2020), Cheating in online courses: Evidence from online proctoring",
            "To study the effect of online proctoring on cheating behavior.",
            "Empirical analysis.",
            "Useful for justifying monitoring and anti-cheat measures.",
            "Focuses on proctoring generally, not on mobile-first CBT design.",
        ],
        [
            "Han and Shin (2016), The Use of a Mobile Learning Management System and Academic Achievement of Online Students",
            "To examine the role of mobile systems in online learning performance.",
            "Study of mobile LMS use and academic performance.",
            "Supports mobile-first educational platform design and access.",
            "Not specifically focused on CBT or anti-cheat controls.",
        ],
    ]
    target_rows = [
        doc.tables[0].rows[1],
        doc.tables[0].rows[2],
        doc.tables[1].rows[0],
        doc.tables[1].rows[1],
        doc.tables[1].rows[2],
    ]
    for row, values in zip(target_rows, literature_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    # Gaps
    set_para_text(
        doc.paragraphs[61],
        "GAPS IDENTIFIED VIS-A-VIS TECHNOLOGY/METHOD TO RESOLVE THE IDENTIFIED GAPS",
    )
    set_para_text(
        doc.paragraphs[62],
        "The reviewed studies show that CBT adoption, online assessment integrity, and mobile learning are all important areas in educational technology. However, most of these works focus on only one part of the problem. Some focus on student perception of CBT, some focus on academic integrity in a broad sense, while others focus on proctoring strategies or mobile learning adoption. Very few bring these concerns together into one department-level, mobile-first CBT system with built-in practical anti-cheat monitoring and live administrative control.",
    )
    set_para_text(
        doc.paragraphs[64],
        "This project addresses that gap by proposing a department-focused CBT platform that combines mobile-first web access, structured exam and question management, timed session control, fullscreen enforcement, tab switch detection, copy and paste blocking, right-click blocking, developer tools detection, warning thresholds, violation logs, and administrative controls such as pause, resume, and final submit. This makes the proposed system more practical for real academic use at department level.",
    )

    # Methodology: remove sample figures
    for index in [77, 76, 74, 73]:
        delete_paragraph(doc.paragraphs[index])

    methodology_heading = find_para(doc, "METHODOLOGY")
    methodology_paragraphs = []
    collecting = False
    for paragraph in doc.paragraphs:
        if paragraph == methodology_heading:
            collecting = True
            continue
        if collecting and paragraph.text.strip() == "EXPECTED RESULT / OUTCOME":
            break
        if collecting and paragraph.text.strip():
            methodology_paragraphs.append(paragraph)

    methodology_texts = [
        "This project will adopt an incremental software development methodology. This approach is suitable because it allows the system to be developed and improved in stages, starting from understanding the problem and ending with testing and evaluation of the completed platform.",
        "The project is also an adapted standalone system. This means that the work will not begin from a completely empty idea. Instead, useful CBT features reviewed from an earlier portal implementation will be studied, restructured, and improved into a focused standalone platform. This makes the project more realistic and gives it a strong technical foundation.",
        "The methodology will cover requirement identification, analysis of the existing system, design of the proposed system, implementation of the standalone product, and testing and evaluation. The main functional areas include student login, exam creation, question management, timing, result handling, session control, and anti-cheat behavior monitoring.",
        "The implementation will use Next.js, React, TypeScript, SQLite or Turso, and Drizzle ORM. Confirmed features from the reviewed technical baseline include exam activation windows, question creation and Excel import, student management, session tracking, question and option shuffling, violation logging, fullscreen enforcement, tab switch detection, copy and paste blocking, right-click blocking, developer tools attempt detection, pause and resume control, and result export.",
        "The completed system will be tested to make sure that students can log in and take exams correctly, administrators can create and manage exams successfully, timers and session controls work correctly, anti-cheat events are logged as expected, and the system remains usable on mobile devices.",
    ]
    while len(methodology_paragraphs) < len(methodology_texts):
        methodology_paragraphs.append(
            insert_paragraph_after(
                methodology_paragraphs[-1] if methodology_paragraphs else methodology_heading,
                style=methodology_heading.style,
            )
        )
    for paragraph, text in zip(methodology_paragraphs, methodology_texts):
        set_para_text(paragraph, text)
    for paragraph in methodology_paragraphs[len(methodology_texts) :]:
        clear_para(paragraph)

    # Expected result
    expected_heading = find_para(doc, "EXPECTED RESULT / OUTCOME")
    expected_heading_index = next(
        i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "EXPECTED RESULT / OUTCOME"
    )
    set_para_text(doc.paragraphs[expected_heading_index + 1], "The proposed study is expected to deliver the following key outcomes:")
    expected_items = [
        "Standalone CBT Platform A mobile-first standalone CBT platform will be developed for the Department of Computer Science, FUPRE.",
        "Administrative Dashboard The system will provide tools for exam setup, student management, question import, session monitoring, and result export.",
        "Student Exam Interface Students will have a secure interface for login, timed testing, answer submission, and result access where permitted.",
        "Practical Anti-Cheat Monitoring The platform will provide warning logs, session monitoring, fullscreen enforcement, and violation-based termination to improve fairness and control during examinations.",
    ]
    expected_list_paragraphs = []
    for paragraph in doc.paragraphs[expected_heading_index + 1 :]:
        if paragraph.style.name == "List Paragraph":
            expected_list_paragraphs.append(paragraph)
            if len(expected_list_paragraphs) == 4:
                break
    for paragraph, text in zip(expected_list_paragraphs, expected_items):
        set_para_text(paragraph, text)

    # Work plan
    work_heading_index = next(
        i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "WORK-PLAN/TIME FRAME USING GANTT CHART"
    )
    set_para_text(
        doc.paragraphs[work_heading_index + 1],
        "This research will span for a period of 3 months spanning from March 2026 to May 2026.",
    )
    work_table = doc.tables[2]
    work_table.rows[0].cells[0].text = "Description"
    work_table.rows[0].cells[1].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[2].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[3].text = "Time Line from March 2026 - May 2026"
    work_table.rows[1].cells[0].text = "Description"
    work_table.rows[1].cells[1].text = "MARCH"
    work_table.rows[1].cells[2].text = "APRIL"
    work_table.rows[1].cells[3].text = "MAY"
    work_activities = [
        ("Study of Existing System/Requirement Gathering", "X", "", ""),
        ("System Analysis and Design", "X", "", ""),
        ("Implementation/Feature Adaptation", "", "X", ""),
        ("Testing, Evaluation, and Report Writing", "", "", "X"),
    ]
    for row_index, values in enumerate(work_activities, start=2):
        for column_index, value in enumerate(values):
            work_table.rows[row_index].cells[column_index].text = value
    set_para_text(
        doc.paragraphs[work_heading_index + 3],
        "The implementation of this project will follow a structured workflow consisting of the following stages:",
    )
    work_item_texts = [
        "Requirement Gathering and Study of Existing System Review of the current CBT process, available technical baseline, and major examination integrity challenges.",
        "System Design Design of the standalone architecture, database structure, user flow, and anti-cheat process.",
        "Implementation Development of the student module, admin module, session control features, and anti-cheat controls.",
        "System Testing Evaluation of functionality, mobile usability, and anti-cheat event handling.",
        "Documentation and Reporting Preparation of the final report, screenshots, findings, and recommendations.",
    ]
    work_list_paragraphs = []
    for paragraph in doc.paragraphs[work_heading_index + 4 :]:
        if paragraph.style.name == "List Paragraph":
            work_list_paragraphs.append(paragraph)
            if len(work_list_paragraphs) == 5:
                break
    for paragraph, text in zip(work_list_paragraphs, work_item_texts):
        set_para_text(paragraph, text)

    # Budget table
    budget_rows = [
        ("Internet/Data", "Data for research, coding, testing, and online access", "25,000"),
        ("Power Support", "Electricity and backup power for development work", "20,000"),
        ("Laptop Maintenance", "Minor maintenance and system support during development", "10,000"),
        ("Transportation", "Movement for project meetings, printing, and logistics", "12,000"),
        ("Printing and Binding", "Draft printing, final printing, and binding", "20,000"),
        ("Hosting/Domain (Optional)", "Optional hosting or domain for system demonstration", "15,000"),
        ("Data Backup/Storage", "Storage support for project files and backups", "8,000"),
        ("Stationery", "Project stationery and related materials", "5,000"),
        ("Testing/Logistics", "Small expenses for user testing and local logistics", "8,000"),
        ("Contingency", "Unexpected expenses", "12,000"),
        ("Total", "", "135,000"),
    ]
    for row, values in zip(doc.tables[3].rows[1:], budget_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    # References
    references_heading_text = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"REFERNCES", "REFERENCES"}:
            references_heading_text = paragraph.text.strip()
            references_heading = paragraph
            break
    if references_heading_text is None:
        raise ValueError("References heading not found")
    set_para_text(references_heading, "REFERENCES")
    reference_texts = [
        "Adamu, A. D. (2024). Computer-Based Testing implementation in Nigeria: Successes and challenges from historical perspectives. UMYU Journal of Educational Research. https://doi.org/10.70886/ujer.24121.013",
        "Akcapinar, G. (2025). Detecting AI-assisted cheating in online exams through behavior analytics. arXiv / CELDA 2025. https://doi.org/10.48550/arXiv.2510.18881",
        "Alessio, H. M., Messinger, J. D., et al. (2021). Faculty and student perceptions of academic integrity in technology-assisted learning and testing. Frontiers in Education, 6. https://doi.org/10.3389/feduc.2021.629220",
        "Cheon, J., Lee, S., Crooks, S. M., & Song, J. (2012). An investigation of mobile learning readiness in higher education based on the theory of planned behavior. Computers & Education, 59(3), 1054-1064. https://doi.org/10.1016/j.compedu.2012.04.015",
        "Chouhan, R. (2023). Strategies for maintaining academic integrity in remote unproctored and proctored online assessments for engineering courses. Learning: Research and Practice, 10(1), 75-92. https://doi.org/10.1080/23735082.2023.2216198",
        "Dendir, S., & Maxwell, R. S. (2020). Cheating in online courses: Evidence from online proctoring. Computers in Human Behavior Reports, 2, 100033. https://doi.org/10.1016/j.chbr.2020.100033",
        "Griffiths, B. J. (2022). Mitigating cheating during online proctored exams. Research on Education and Media, 14(2), 9-14. https://doi.org/10.2478/rem-2022-0016",
        "Han, I., & Shin, W. S. (2016). The use of a mobile learning management system and academic achievement of online students. Computers & Education, 102, 79-89. https://doi.org/10.1016/j.compedu.2016.07.003",
        "Holden, O. L., Norris, M. E., & Kuhlmeier, V. A. (2021). Academic integrity in online assessment: A research review. Frontiers in Education, 6, 639814. https://doi.org/10.3389/feduc.2021.639814",
        "Naveed, Q. N., Choudhary, H., Ahmad, N., Alqahtani, J., & Qahmash, A. I. (2023). Mobile learning in higher education: A systematic literature review. Sustainability, 15(18), 13566. https://doi.org/10.3390/su151813566",
        "Okocha, F. (2022). Student perception of computer-based testing in Kwara State, Nigeria. International Journal of Web-Based Learning and Teaching Technologies. https://doi.org/10.4018/IJWLTT.294575",
        "Tella, A., & Bashorun, M. T. (2012). Attitude of undergraduate students towards computer-based test (CBT): A case study of the University of Ilorin, Nigeria. International Journal of Information and Communication Technology Education. https://doi.org/10.4018/jicte.2012040103",
    ]
    while True:
        current_paragraphs = doc.paragraphs
        heading_index = next(
            (i for i, paragraph in enumerate(current_paragraphs) if paragraph.text.strip() == "REFERENCES"),
            -1,
        )
        if heading_index == -1:
            raise ValueError("Updated references heading not found")
        if heading_index + 1 >= len(current_paragraphs):
            break
        delete_paragraph(current_paragraphs[heading_index + 1])
    last_reference = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "REFERENCES")
    for reference in reference_texts:
        last_reference = insert_paragraph_after(last_reference, text=reference, style=references_heading.style)

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")
    print(f"Backup saved at {BACKUP_PATH}")


if __name__ == "__main__":
    main()
