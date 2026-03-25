from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TOPIC_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = TOPIC_DIR / "drafts" / "Preamble - Azeta Duke.docx"

TITLE = (
    "DESIGN AND IMPLEMENTATION OF A MOBILE-FIRST ANTI-CHEAT CBT PLATFORM "
    "FOR THE DEPARTMENT OF COMPUTER SCIENCE, FEDERAL UNIVERSITY OF PETROLEUM "
    "RESOURCES, EFFURUN"
)


def set_default_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_page_number_footer(section, *, roman: bool, hide_first: bool) -> None:
    section.different_first_page_header_footer = hide_first

    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    pg_num.set(qn("w:start"), "1")
    pg_num.set(qn("w:fmt"), "lowerRoman" if roman else "decimal")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    bold: bool = False,
    size: int = 12,
    uppercase: bool = False,
    first_line_indent: float = 0.0,
    space_after: int = 8,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.line_spacing = 2.0

    run = paragraph.add_run(text.upper() if uppercase else text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_blank_lines(doc: Document, count: int) -> None:
    for _ in range(count):
        add_paragraph(doc, "", space_after=0)


def add_page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_signature_block(doc: Document, name: str, role: str) -> None:
    add_paragraph(doc, "______________________________", space_after=2)
    add_paragraph(doc, name, space_after=2)
    add_paragraph(doc, role, space_after=2)
    add_paragraph(doc, "Date: ______________________", space_after=10)


def build_cover_page(doc: Document) -> None:
    add_blank_lines(doc, 2)
    add_paragraph(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "[UNIVERSITY LOGO]", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "AZETA DUKE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 1)
    add_paragraph(doc, "COS/8650/2021", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "DEPARTMENT OF COMPUTER SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 1)
    add_paragraph(doc, "COLLEGE OF SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 1)
    add_paragraph(
        doc,
        "FEDERAL UNIVERSITY OF PETROLEUM RESOURCES, EFFURUN, DELTA STATE, NIGERIA",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        uppercase=True,
    )
    add_blank_lines(doc, 3)
    add_paragraph(doc, "MARCH, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)


def build_title_page(doc: Document) -> None:
    add_blank_lines(doc, 2)
    add_paragraph(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 1)
    add_paragraph(doc, "AZETA DUKE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 1)
    add_paragraph(doc, "COS/8650/2021", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)
    add_blank_lines(doc, 2)
    add_paragraph(
        doc,
        "A PROJECT SUBMITTED TO THE",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        uppercase=True,
    )
    add_paragraph(
        doc,
        "DEPARTMENT OF COMPUTER SCIENCE,",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        uppercase=True,
    )
    add_paragraph(
        doc,
        "COLLEGE OF SCIENCE, FEDERAL UNIVERSITY OF PETROLEUM RESOURCES,",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        uppercase=True,
    )
    add_paragraph(
        doc,
        "EFFURUN, DELTA STATE",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        uppercase=True,
    )
    add_blank_lines(doc, 1)
    add_paragraph(
        doc,
        "IN PARTIAL FULFILMENT FOR THE AWARD OF A BACHELOR OF",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        uppercase=True,
    )
    add_paragraph(
        doc,
        "SCIENCE (B.Sc.) DEGREE IN COMPUTER SCIENCE",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=False,
        uppercase=True,
    )
    add_blank_lines(doc, 3)
    add_paragraph(doc, "MARCH, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, uppercase=True)


def build_declaration(doc: Document) -> None:
    add_paragraph(doc, "DECLARATION", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 1)
    add_paragraph(
        doc,
        'I, Azeta Duke, with matriculation number COS/8650/2021, hereby declare that this project '
        'titled, "Design and Implementation of a Mobile-First Anti-Cheat CBT Platform for the '
        'Department of Computer Science, Federal University of Petroleum Resources, Effurun" was '
        "carried out by me under the supervision of Dr Nwozor Blessing, in partial fulfilment of "
        "the requirement for the award of B.Sc. Computer Science from the Department of Computer "
        "Science, Federal University of Petroleum Resources, Effurun. This project is original and "
        "has not been submitted in part or in full for the award of any degree in this or any "
        "other institution.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=0.5,
    )
    add_blank_lines(doc, 2)
    add_signature_block(doc, "Azeta Duke", "Student")
    add_blank_lines(doc, 1)
    add_signature_block(doc, "Dr Nwozor Blessing", "Supervisor")


def build_certification(doc: Document) -> None:
    add_paragraph(doc, "CERTIFICATION", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 1)
    add_paragraph(
        doc,
        'This is to certify that this project titled, "Design and Implementation of a Mobile-First '
        'Anti-Cheat CBT Platform for the Department of Computer Science, Federal University of '
        "Petroleum Resources, Effurun\" was carried out by Azeta Duke with matriculation number "
        "COS/8650/2021, and has been approved by the undersigned having met the partial requirement "
        "for the award of a Bachelor of Science (B.Sc.) degree in Computer Science from the "
        "Department of Computer Science, Federal University of Petroleum Resources, Effurun.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=0.5,
    )
    add_blank_lines(doc, 2)
    add_signature_block(doc, "Dr Nwozor Blessing", "Project Supervisor")
    add_blank_lines(doc, 1)
    add_signature_block(doc, "Prof./Dr. ______________________", "Head of Department")
    add_blank_lines(doc, 1)
    add_signature_block(doc, "______________________", "External Examiner")


def build_dedication(doc: Document) -> None:
    add_paragraph(doc, "DEDICATION", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 2)
    add_paragraph(
        doc,
        "This project is dedicated to God Almighty for His grace, strength, and guidance throughout "
        "the course of this study. It is also dedicated to my family for their support, "
        "encouragement, and prayers.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=0.5,
    )


def build_acknowledgements(doc: Document) -> None:
    add_paragraph(doc, "ACKNOWLEDGEMENTS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 1)
    paragraphs = [
        "I give thanks to God Almighty for His mercy, wisdom, and strength throughout the period of this project work and my academic programme.",
        "I sincerely appreciate my supervisor, Dr Nwozor Blessing, for the guidance, correction, encouragement, and support given to me during the course of this work. Your advice and supervision were very helpful to the success of this project.",
        "I also appreciate the lecturers and staff of the Department of Computer Science, Federal University of Petroleum Resources, Effurun, for the knowledge and support they provided during my period of study.",
        "My sincere gratitude goes to my parents, family members, and friends for their encouragement, prayers, and understanding throughout this project.",
        "I am grateful to everyone who contributed in one way or another to the success of this study.",
    ]
    for text in paragraphs:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=0.5)


def build_abstract(doc: Document) -> None:
    add_paragraph(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 1)
    paragraphs = [
        "This study focused on the design and implementation of a mobile-first anti-cheat Computer-Based Testing platform for the Department of Computer Science, Federal University of Petroleum Resources, Effurun. The aim of the study was to develop a standalone CBT system that improved examination access on mobile devices and supported practical anti-cheat measures during examinations.",
        "The system was designed as a web-based platform with separate interfaces for administrators and students. The administrator section was used for exam creation, question management, student management, session monitoring, and result handling, while the student section was used for login, timed examination access, answer submission, and result viewing where permitted. Practical anti-cheat features such as fullscreen monitoring, tab switch detection, copy and paste restriction, right-click restriction, question randomization, and violation logging were implemented in the system.",
        "The system was developed to improve examination management, fairness, and usability within the department. The result showed that the platform provided a more organized way to conduct computer-based examinations while also supporting better monitoring of student sessions. It was concluded that a mobile-first anti-cheat CBT platform could improve digital examination practice at the department level.",
    ]
    for text in paragraphs:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=0.5)


def build_placeholder_page(doc: Document, title: str, body_lines: list[str]) -> None:
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_blank_lines(doc, 1)
    for line in body_lines:
        add_paragraph(doc, line, first_line_indent=0.0)


def main() -> None:
    doc = Document()
    set_default_font(doc)
    set_margins(doc)
    add_page_number_footer(doc.sections[0], roman=True, hide_first=True)

    build_cover_page(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_title_page(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_declaration(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_certification(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_dedication(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_acknowledgements(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_abstract(doc)
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_placeholder_page(
        doc,
        "TABLE OF CONTENTS",
        [
            "Declaration",
            "Certification",
            "Dedication",
            "Acknowledgements",
            "Abstract",
            "Table of Contents",
            "List of Figures",
            "List of Tables",
            "CHAPTER ONE",
            "INTRODUCTION",
            "1.1 BACKGROUND OF THE STUDY",
            "1.2 STATEMENT OF THE PROBLEM",
            "1.3 AIM AND OBJECTIVES OF THE STUDY",
            "1.4 SIGNIFICANCE OF THE STUDY",
            "1.5 SCOPE OF THE STUDY",
            "1.6 DEFINITION OF TERMS",
            "CHAPTER TWO",
            "LITERATURE REVIEW",
            "2.1 THEORETICAL REVIEW OR CONCEPTUAL REVIEW",
            "2.2 OTHER LITERATURE ON PREVIOUS APPROACHES, THEORIES, AND TECHNIQUES",
            "2.3 REVIEW OF RELATED WORKS",
            "2.4 SUMMARY OF LITERATURE REVIEW",
            "2.5 GAPS IN RELATED WORKS",
            "CHAPTER THREE",
            "CHAPTER FOUR",
            "CHAPTER FIVE",
            "References",
            "Appendices",
        ],
    )
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_placeholder_page(
        doc,
        "LIST OF FIGURES",
        [
            "Figure 3.1: System Architecture",
            "Figure 3.2: Use Case Diagram",
            "Figure 4.1: Admin Dashboard",
            "Figure 4.2: Student Exam Interface",
        ],
    )
    add_page_break(doc)
    add_page_number_footer(doc.sections[-1], roman=True, hide_first=False)
    build_placeholder_page(
        doc,
        "LIST OF TABLES",
        [
            "Table 2.1: Summary of Reviewed Literature",
            "Table 3.1: System Requirements",
            "Table 4.1: Test Cases and Results",
        ],
    )

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
