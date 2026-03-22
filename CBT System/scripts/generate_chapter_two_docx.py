from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TOPIC_DIR = Path(__file__).resolve().parents[1]
INPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 02 Draft.md"
OUTPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 02 - Azeta Duke.docx"


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.space_after = Pt(4)
    add_markdown_runs(paragraph, text, 12)


def add_list_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    add_markdown_runs(paragraph, text, 12)


def add_markdown_runs(paragraph, text: str, size: int, *, bold_default: bool = False) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue

        bold = bold_default
        italic = False
        content = part

        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            bold = True
        elif part.startswith("*") and part.endswith("*"):
            content = part[1:-1]
            italic = True

        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)


def split_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        for line in table_lines:
            add_body_paragraph(doc, line)
        return

    headers = split_markdown_row(table_lines[0])
    data_rows = [split_markdown_row(line) for line in table_lines[2:]]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        paragraph = table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_markdown_runs(paragraph, header, 11, bold_default=True)

    for row_values in data_rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            paragraph = row.cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_markdown_runs(paragraph, value, 11)


def main() -> None:
    doc = Document()
    set_default_font(doc)
    set_margins(doc)

    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
        elif line[:2].isdigit() and line[1] == ".":
            add_list_paragraph(doc, line)
        elif line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines):
                current = lines[i].strip()
                if current.startswith("|") and current.endswith("|"):
                    table_lines.append(current)
                    i += 1
                else:
                    break
            add_markdown_table(doc, table_lines)
            continue
        else:
            add_body_paragraph(doc, line)
        i += 1

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
