from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE = Path(r"C:\Users\azeta\Documents\School\Project Topic")
TEMPLATE_PATH = BASE / "Project Proposal.original.docx"
OUTPUT_PATH = BASE / "Project Proposal - Azeta Duke - Pharmacy Inventory.docx"


def has_drawing(run) -> bool:
    xml = run._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def set_para_text(paragraph, text: str) -> None:
    text_runs = [run for run in paragraph.runs if not has_drawing(run)]
    if text_runs:
        text_runs[0].text = text
        for run in text_runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_para(paragraph) -> None:
    set_para_text(paragraph, "")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def collect_between(doc: Document, start_text: str, end_text: str):
    start = find_para(doc, start_text)
    items = []
    collecting = False
    for paragraph in doc.paragraphs:
        if paragraph == start:
            collecting = True
            continue
        if collecting and paragraph.text.strip() == end_text:
            break
        if collecting:
            items.append(paragraph)
    return items


def main() -> None:
    doc = Document(TEMPLATE_PATH)

    title = (
        "Project Title: Design and Implementation of a Mobile Pharmacy Inventory and "
        "Expiry Tracking System for Community Pharmacies in Effurun"
    )
    executive_summary = [
        "Community pharmacies play an important role in medicine access, but many still manage stock manually or with weak record systems. This can make it difficult to know the real stock level of medicines, track batch details, monitor expiry dates, and respond quickly to low-stock situations. When inventory control is weak, the result may be medicine stockouts, expired products on the shelf, product waste, and poor service delivery.",
        "This project focuses on the design and implementation of a mobile pharmacy inventory and expiry tracking system for community pharmacies in Effurun. The system will be built as a standalone mobile-first web application that can be used easily on smartphones and also on laptops where available. It will allow pharmacy staff to register medicines, record batch details, track stock in and stock out, monitor expiry dates, identify low-stock items, and view simple inventory reports through a dashboard. The proposed system is not intended to be a full pharmacy point-of-sale platform. Its main goal is to improve medicine stock control and expiry monitoring in community pharmacies. The expected result is a practical inventory support tool that helps reduce stockouts, reduce waste from expired medicines, and improve the everyday inventory workflow of community pharmacies in Effurun.",
    ]
    background = [
        "Medicine availability is an important part of quality healthcare. When essential medicines are unavailable, patients may not get treatment on time, and pharmacy operations become less reliable. The Federal Ministry of Health and Social Welfare recently launched a digital inventory model for essential medicines as part of efforts to reduce drug stockouts in Nigeria. This shows that inventory visibility and stock control are now receiving greater attention in the Nigerian health sector (Federal Ministry of Health and Social Welfare, 2025).",
        "Community pharmacies are often the first point of contact for medicine access in many communities. For these pharmacies to function well, staff must be able to know what medicines are available, what quantity remains, what batch was received, and which products are close to expiry. However, when records are kept manually, this process can become slow and error-prone. Weak inventory practice may also lead to poor medicine disposal decisions and higher losses.",
        "Studies from Nigeria show that expired and unused medicine management remains an important issue in pharmacy practice. Iweh et al. (2019) examined disposal practices among community pharmacies in southeast Nigeria and showed that expired medication management is a real concern in local pharmacy settings. Akande-Sholabi et al. (2025) also reported results on unused and expired medication practices among healthcare practitioners in Ibadan. These findings support the need for better systems that make medicine tracking easier and more accurate before disposal even becomes necessary.",
        "Inventory control methods have also been studied in pharmacy and medicine supply settings beyond Nigeria. Watson et al. (2014) examined inventory control methods in a pharmacy environment and showed the importance of structured inventory practices. Sukendar et al. (2020) further showed that medicine inventory control should consider expiry periods and product returns. These ideas are relevant to this study because a community pharmacy inventory system should not only count products, but should also track expiry risk and give timely alerts. Based on this background, there is a clear need for a mobile-first digital system that supports community pharmacy staff in Effurun with stock visibility, batch tracking, expiry tracking, and simple reporting.",
    ]
    problem = [
        "Many community pharmacies still rely on manual record keeping or simple methods that do not give a clear real-time view of medicine stock. This makes it difficult for pharmacy staff to know what medicines are available, which ones are low in quantity, which batches are nearing expiry, and which products have already expired. As a result, pharmacies may experience stockouts, delayed restocking, poor inventory decisions, and medicine waste.",
        "Another challenge is that pharmacy workers may need to update stock records while moving around the shop, not only from a desktop computer. A system that is not mobile-friendly may reduce ease of use and discourage regular updates. If stock information is not entered quickly and accurately, the quality of inventory records becomes weaker.",
        "The problem this project addresses, therefore, is the absence of a mobile-first inventory and expiry tracking system designed specifically for community pharmacies in Effurun. Such a system is needed to support medicine registration, batch-based stock tracking, low-stock alerts, near-expiry alerts, and simple reporting in a more practical and reliable way.",
    ]
    objectives = [
        "To build a mobile-first web system for inventory management in community pharmacies.",
        "To provide support for medicine registration, batch entry, and quantity tracking.",
        "To implement expiry date tracking for medicine batches.",
        "To provide low-stock and near-expiry dashboard alerts.",
        "To support stock-in and stock-out recording.",
        "To provide simple inventory reports for pharmacy staff.",
    ]
    conceptual = [
        "Pharmacy inventory management refers to the process of recording, monitoring, and controlling the movement of medicines and related products in a pharmacy. A good inventory system helps ensure that medicines are available when needed and that stock records remain accurate.",
        "Expiry tracking is the process of monitoring the expiry dates of medicine batches so that products close to expiry can be identified early. This is important because expired medicines should not remain in active stock for patient use. In this project, expiry tracking is one of the central functions of the system.",
        "A low-stock alert system helps users know when the quantity of a product has fallen below a defined level. In a pharmacy setting, this supports timely restocking and helps reduce the risk of stockouts. In this project, alerts will appear on the dashboard rather than through SMS or email.",
        "Batch tracking means recording each medicine batch separately with details such as batch number, quantity, and expiry date. This is important in pharmacy operations because the same medicine may exist in different batches with different expiry dates.",
        "Community pharmacy workflow includes medicine receiving, stock recording, storage, stock movement, expiry monitoring, and reporting. This project is designed specifically around that workflow, not around hospital pharmacy or full point-of-sale operations.",
    ]
    literature_rows = [
        ["Iweh et al. (2019), Assessment of disposal practices of expired and unused medications among community pharmacies in Anambra State southeast Nigeria", "To examine disposal practices of expired and unused medicines in community pharmacies.", "Mixed study design.", "Gives direct Nigerian community-pharmacy evidence that expired medicine handling is a real issue.", "Focuses on disposal practice, not on building a digital stock and expiry system."],
        ["Akande-Sholabi et al. (2025), Disposal practices of unused and expired medications among healthcare practitioners in Ibadan, Nigeria", "To study expired and unused medicine practices among healthcare practitioners.", "Cross-sectional survey.", "Provides recent Nigerian evidence that expired medicine management remains important.", "Does not focus specifically on inventory software for community pharmacies."],
        ["Watson et al. (2014), Inventory Control Methods in a Long-Term Care Pharmacy: Comparisons and Time-Series Analyses", "To compare inventory control methods in a pharmacy setting.", "Comparative and time-series analysis.", "Shows the value of structured inventory control in pharmacy practice.", "Not focused on mobile-first systems or community pharmacies in Nigeria."],
        ["Sukendar et al. (2020), Medicine Inventory Control by Considering Expiry Periods and Product Returns...", "To improve medicine inventory control while considering expiry periods and returns.", "Inventory control modelling and analysis.", "Highlights the importance of expiry-aware inventory management.", "Not tailored to Nigerian community pharmacy workflow."],
        ["Federal Ministry of Health and Social Welfare (2025), FG Moves to Eliminate Drug Stockouts, Launches Digital Inventory Model for Essential Medicines", "To improve medicine availability and reduce stockouts through digital inventory control.", "National policy and implementation direction.", "Shows that digital inventory control is currently relevant in Nigeria's health sector.", "Policy-level direction; does not provide a focused community pharmacy software solution."],
    ]
    methodology = [
        "This project will adopt an incremental software development methodology. This approach is suitable because it allows the system to be developed in stages, tested gradually, and improved as each major part is completed.",
        "The methodology will begin with requirement gathering and study of inventory workflow in community pharmacies. This will involve identifying the main users of the system, the key stock records required, and the major problems related to stockouts, weak stock visibility, and expiry management.",
        "The next stage will involve system analysis and design. At this stage, the database structure, user interface, dashboard flow, and alert logic will be planned. The proposed system will include medicine registration, batch entry, quantity tracking, stock movement records, low-stock alerts, near-expiry alerts, and simple inventory reports.",
        "The implementation stage will focus on building the mobile-first web application and connecting it to a database for storing medicine and batch records. The prototype can be developed using tools such as Next.js, React, TypeScript, and a relational database such as SQLite or MySQL.",
        "The final stage will be testing and evaluation. The system will be tested to confirm that medicines can be added correctly, quantities can be updated correctly, expiry dates can be tracked, alerts are shown properly, and the dashboard remains easy to use on both smartphones and laptops.",
    ]
    expected_items = [
        "Mobile Pharmacy Inventory Dashboard A mobile-first dashboard for community pharmacy staff to manage medicine stock records.",
        "Expiry Tracking Module A system component that records batch expiry dates and flags medicines that are close to expiry or already expired.",
        "Low-Stock Alert Support A dashboard view that highlights medicines that have dropped below the required stock level.",
        "Improved Inventory Control A practical digital tool that can help reduce stockouts, reduce waste, and improve everyday stock management in community pharmacies in Effurun.",
    ]
    work_items = [
        "Requirement Gathering and Workflow Study Review of community pharmacy inventory workflow, stock control problems, and expiry management needs.",
        "System Analysis and Design Design of the database, mobile-first interface, and alert dashboard.",
        "Implementation Development of medicine registration, batch entry, stock tracking, expiry tracking, and reporting features.",
        "System Testing Evaluation of stock updates, expiry alerts, low-stock alerts, and mobile usability.",
        "Documentation and Reporting Preparation of final report, screenshots, findings, and recommendations.",
    ]
    budget_rows = [
        ("Internet/Data", "Data for research, development, testing, and online access", "20,000"),
        ("Power Support", "Electricity and backup power during development", "20,000"),
        ("Transportation", "Visits for project meetings, logistics, and printing", "15,000"),
        ("Printing and Binding", "Draft printing, final printing, and binding", "20,000"),
        ("Local Testing/Field Visits", "Small expenses for pharmacy visits and workflow observation", "10,000"),
        ("Hosting/Domain (Optional)", "Optional hosting or domain for demonstration", "15,000"),
        ("Data Backup/Storage", "Backup support for project files and records", "8,000"),
        ("Stationery", "Project stationery and related materials", "5,000"),
        ("User Testing/Logistics", "Small expenses for local testing and coordination", "10,000"),
        ("Contingency", "Unexpected expenses", "12,000"),
        ("Total", "", "135,000"),
    ]
    references = [
        "Akande-Sholabi, W., Abdul-Azeez, I. A., Adebisi, Y. A., Odukoya, T. O., & Ilori, T. (2025). Disposal practices of unused and expired medications among healthcare practitioners in Ibadan, Nigeria: Results from a cross-sectional survey. BMC Health Services Research, 25, 1262. https://doi.org/10.1186/s12913-025-13492-0",
        "Federal Ministry of Health and Social Welfare. (2025). FG moves to eliminate drug stockouts, launches digital inventory model for essential medicines. https://health.gov.ng/fg-moves-to-eliminate-drug-stockouts-launches-digital-inventory-model-for-essential-medicines/",
        "Federal Government of Nigeria. (2018). Second National Strategic Health Development Plan 2018-2022. https://extranet.who.int/countryplanningcycles/sites/default/files/public_file_rep/NGA_Nigeria_Second-National-Strategic-Health-Development-Plan_2018-2022.pdf",
        "Iweh, M., Ogbonna, B., Sunday, N., Anetoh, M., & Okonta, M. (2019). Assessment of disposal practices of expired and unused medications among community pharmacies in Anambra State southeast Nigeria: A mixed study design. Journal of Pharmaceutical Policy and Practice, 12, 12. https://doi.org/10.1186/s40545-019-0174-1",
        "Sukendar, I., Sugiyono, A., & Munfiqotusshifa. (2020). Medicine inventory control by considering expiry periods and product returns using the always better control (ABC) analysis and the Handley within model of economic order quality (EOQ) at pharmacies in Indonesia. Journal of Technology and Operations Management, 15(2), 20-30. https://doi.org/10.32890/jtom2020.15.2.3",
        "Watson, J. W., Moliver, N., & Gossett, K. (2014). Inventory control methods in a long-term care pharmacy: Comparisons and time-series analyses. Journal of Pharmacy Technology, 30(5), 151-162. https://doi.org/10.1177/8755122514534073",
    ]

    set_para_text(doc.paragraphs[5], "Name: AZETA DUKE")
    set_para_text(doc.paragraphs[6], "Matric Number: COS/8650/2021 Supervisor/Mentor: Dr NWOZOR BLESSING Department: COMPUTER SCIENCE")
    set_para_text(doc.paragraphs[7], "Institution: FEDERAL UNIVERSITY OF PETROLEUM RESOURCES Academic Session: 2025/2026")
    set_para_text(doc.paragraphs[10], title)

    set_para_text(doc.paragraphs[14], executive_summary[0])
    set_para_text(doc.paragraphs[16], executive_summary[1])

    for index, text in zip([19, 21, 23, 25], background):
        set_para_text(doc.paragraphs[index], text)
    clear_para(doc.paragraphs[27])
    clear_para(doc.paragraphs[29])

    for index, text in zip([32, 34, 35], problem):
        set_para_text(doc.paragraphs[index], text)

    set_para_text(doc.paragraphs[38], "To design and implement a mobile pharmacy inventory and expiry tracking system for community pharmacies in Effurun. The specific objectives of the study are as follows:")

    for index, text in zip([44, 46, 48, 50, 52], conceptual):
        set_para_text(doc.paragraphs[index], text)
    for index in [45, 47, 49, 51, 53, 54]:
        if index < len(doc.paragraphs):
            clear_para(doc.paragraphs[index])

    target_rows = [doc.tables[0].rows[1], doc.tables[0].rows[2], doc.tables[1].rows[0], doc.tables[1].rows[1], doc.tables[1].rows[2]]
    for row, values in zip(target_rows, literature_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    set_para_text(doc.paragraphs[62], "The reviewed works show that medicine stock control, expiry management, and disposal practices remain important pharmacy issues. Some studies focus on disposal practices, some focus on inventory control methods, and others show policy interest in digital inventory solutions. However, very few of them propose a mobile-first inventory and expiry tracking system designed specifically for community pharmacies in Effurun.")
    set_para_text(doc.paragraphs[64], "This project addresses that gap by proposing a mobile pharmacy inventory and expiry tracking system that combines medicine registration, batch entry, stock-in and stock-out tracking, quantity monitoring, low-stock alerts, near-expiry alerts, and simple reports in one mobile-friendly platform. This makes the proposed system more practical for community pharmacies in Effurun where staff need a simple way to manage stock and reduce medicine waste.")

    section_paragraphs = collect_between(doc, "AIM/OBJECTIVE", "CONCEPTUAL REVIEW")
    objective_paragraphs = [p for p in section_paragraphs if p.style.name == "List Paragraph"]
    anchor = objective_paragraphs[-1] if objective_paragraphs else doc.paragraphs[38]
    while len(objective_paragraphs) < len(objectives):
        anchor = insert_paragraph_after(anchor, style=doc.paragraphs[39].style)
        objective_paragraphs.append(anchor)
    for paragraph, text in zip(objective_paragraphs[: len(objectives)], objectives):
        set_para_text(paragraph, text)
    for paragraph in objective_paragraphs[len(objectives) :]:
        delete_paragraph(paragraph)

    methodology_heading = find_para(doc, "METHODOLOGY")
    methodology_section = collect_between(doc, "METHODOLOGY", "EXPECTED RESULT / OUTCOME")
    methodology_style = None
    for paragraph in methodology_section:
        if paragraph.text.strip():
            methodology_style = paragraph.style
            break
    if methodology_style is None:
        methodology_style = doc.paragraphs[67].style
    for paragraph in reversed(methodology_section):
        delete_paragraph(paragraph)
    last_methodology = methodology_heading
    for text in methodology:
        last_methodology = insert_paragraph_after(last_methodology, text=text, style=methodology_style)

    expected_heading_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "EXPECTED RESULT / OUTCOME")
    set_para_text(doc.paragraphs[expected_heading_index + 1], "The proposed study is expected to deliver the following key outcomes:")
    expected_list_paragraphs = []
    for paragraph in doc.paragraphs[expected_heading_index + 1 :]:
        if paragraph.style.name == "List Paragraph":
            expected_list_paragraphs.append(paragraph)
            if len(expected_list_paragraphs) == 4:
                break
    for paragraph, text in zip(expected_list_paragraphs, expected_items):
        set_para_text(paragraph, text)

    work_heading_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "WORK-PLAN/TIME FRAME USING GANTT CHART")
    set_para_text(doc.paragraphs[work_heading_index + 1], "This research will span for a period of 3 months spanning from March 2026 to May 2026.")
    work_table = doc.tables[2]
    work_table.rows[0].cells[0].text = "Description"
    work_table.rows[0].cells[1].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[2].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[3].text = "Time Line from March 2026 - May 2026"
    work_table.rows[1].cells[0].text = "Description"
    work_table.rows[1].cells[1].text = "MARCH"
    work_table.rows[1].cells[2].text = "APRIL"
    work_table.rows[1].cells[3].text = "MAY"
    work_activities = [("Requirement Gathering/Workflow Study", "X", "", ""), ("System Analysis and Design", "X", "", ""), ("Implementation", "", "X", ""), ("Testing, Evaluation, and Report Writing", "", "", "X")]
    for row_index, values in enumerate(work_activities, start=2):
        for column_index, value in enumerate(values):
            work_table.rows[row_index].cells[column_index].text = value
    set_para_text(doc.paragraphs[work_heading_index + 3], "The implementation of this project will follow a structured workflow consisting of the following stages:")
    work_list_paragraphs = []
    for paragraph in doc.paragraphs[work_heading_index + 4 :]:
        if paragraph.style.name == "List Paragraph":
            work_list_paragraphs.append(paragraph)
            if len(work_list_paragraphs) == 5:
                break
    for paragraph, text in zip(work_list_paragraphs, work_items):
        set_para_text(paragraph, text)
    intro_text = "The implementation of this project will follow a structured workflow consisting of the following stages:"

    for row, values in zip(doc.tables[3].rows[1:], budget_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    references_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"REFERNCES", "REFERENCES"}:
            references_heading = paragraph
            break
    if references_heading is None:
        raise ValueError("References heading not found")
    set_para_text(references_heading, "REFERENCES")
    while True:
        current_paragraphs = doc.paragraphs
        heading_index = next((i for i, paragraph in enumerate(current_paragraphs) if paragraph.text.strip() == "REFERENCES"), -1)
        if heading_index == -1 or heading_index + 1 >= len(current_paragraphs):
            break
        delete_paragraph(current_paragraphs[heading_index + 1])
    last_reference = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "REFERENCES")
    for reference in references:
        last_reference = insert_paragraph_after(last_reference, text=reference, style=references_heading.style)

    intro_hits = [p for p in doc.paragraphs if p.text.strip() == intro_text]
    for paragraph in reversed(intro_hits[1:]):
        delete_paragraph(paragraph)

    banned_phrases = [
        "parallel bilingual corpus",
        "pre-trained T5 transformer",
        "Englishâ€“Itsekiri",
        "English–Itsekiri translation",
        "trained model",
        "Figure 1: Translation process overview",
        "Figure 2: The Architecture",
        "For the purpose of this work",
        "T5 functions as a sequence-to-sequence model",
        "Upon completion of training and fine-tuning",
        "Through this platform, users will input English text",
    ]
    for paragraph in reversed(list(doc.paragraphs)):
        text = paragraph.text.strip()
        if any(phrase in text for phrase in banned_phrases):
            delete_paragraph(paragraph)

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
