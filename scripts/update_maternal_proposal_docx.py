from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE = Path(r"C:\Users\azeta\Documents\School\Project Topic")
TEMPLATE_PATH = BASE / "Project Proposal.original.docx"
OUTPUT_PATH = BASE / "Project Proposal - Azeta Duke - PHC Maternal Care.docx"


def has_drawing(run) -> bool:
    xml = run._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def set_para_text(paragraph, text: str) -> None:
    text_runs = [run for run in paragraph.runs if not has_drawing(run)]
    if text_runs:
        text_runs[0].text = text
        for run in text_runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_para(paragraph) -> None:
    set_para_text(paragraph, "")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def collect_between(doc: Document, start_text: str, end_text: str):
    start = find_para(doc, start_text)
    items = []
    collecting = False
    for paragraph in doc.paragraphs:
        if paragraph == start:
            collecting = True
            continue
        if collecting and paragraph.text.strip() == end_text:
            break
        if collecting:
            items.append(paragraph)
    return items


def main() -> None:
    doc = Document(TEMPLATE_PATH)

    title = (
        "Project Title: Design and Implementation of a Mobile-First SMS-Based Maternal Care "
        "Reminder and Follow-Up Platform for Government Primary Health Centres in Effurun"
    )
    executive_summary = [
        "Maternal health remains a serious public health concern in Nigeria, and many pregnant women still miss important antenatal and postnatal visits for different reasons. In many government Primary Health Centres (PHCs), follow-up is still done manually, records may be paper-based, and health workers often have limited access to standard computers. At the same time, many women may not own smartphones or may not be confident using digital applications by themselves. These realities make it necessary to design a system that is simple, inclusive, and practical for the local environment.",
        "This project focuses on the design and implementation of a mobile-first SMS-based maternal care reminder and follow-up platform for government Primary Health Centres in Effurun. The system will be built as a standalone web platform that works well on smartphones and can also be accessed on PCs when available. Health workers will use the system to register pregnant women, record antenatal and postnatal visit dates, track women due for checkups, identify missed visits, and manage follow-up actions. On the patient side, women will not need an app. They will receive simple SMS reminders before scheduled visits and follow-up SMS messages after missed appointments. The proposed system is designed to fit the realities of government PHCs where staff may depend more on smartphones than desktop systems, and where patients may use either basic phones or smartphones. The expected result is a mobile-first clinic support platform that improves appointment attendance, strengthens follow-up, reduces missed maternal visits, and supports better record handling and monitoring in government PHCs in Effurun.",
    ]
    background = [
        "Maternal and newborn health remains an important issue in Nigeria. UNICEF notes that Nigeria carries a heavy share of global maternal deaths, and this shows the need for stronger support systems that help women stay connected to care before and after delivery (UNICEF, n.d.). Regular antenatal and postnatal visits are important because they help health workers detect risks early, monitor pregnancy progress, and provide support for both mother and child. However, the use of maternal healthcare services is still uneven across the country, and this affects the quality of outcomes for women and babies (Oyinlola et al., 2025).",
        "Government Primary Health Centres are often the first point of contact for maternal care at community level. They play an important role in registration, routine checks, follow-up, and referral when necessary. Even so, the quality of maternal and newborn care in many primary health facilities still needs improvement, especially in the areas of continuity of care, follow-up, and access to timely services (Oluwatola et al., 2025). When women miss scheduled visits and there is no proper follow-up system, the risk of poor maternal outcomes can increase.",
        "Digital health tools have created opportunities to improve communication between health facilities and patients. Mobile health approaches are especially useful in environments where mobile phones are more common than computers. Research has shown that reminder systems based on SMS and voice messaging can improve attendance and continuity of maternal care. For example, Osanyin et al. (2022) reported positive results from a voice messaging intervention for antenatal care in Lagos, while Olajubu et al. (2020) showed that a mobile health intervention improved uptake of postnatal care services in Nigeria. Broader review studies also show that SMS reminders can improve antenatal care and related maternal service use in developing countries (Hailemariam et al., 2024; Kante & Malqvist, 2025).",
        "Despite these opportunities, many women still face barriers such as low digital literacy, limited internet access, and dependence on basic mobile phones. For this reason, a maternal reminder system should not depend only on smartphone apps. It should be simple enough for health workers to operate and flexible enough to reach pregnant women through basic SMS communication. This project addresses that need by proposing a mobile-first web platform for health workers in government PHCs in Effurun, supported by SMS reminders for pregnant women.",
    ]
    problem = [
        "In many government Primary Health Centres, maternal appointment tracking and follow-up are still handled manually. This can make it difficult for health workers to know which women are due for antenatal or postnatal visits, who has missed appointments, and who needs immediate follow-up. When records are not easy to access or update, continuity of care becomes weaker and important visits may be missed.",
        "Another challenge is that many PHC workers may not always have access to standard computers during their daily work, even though smartphones are more available. Also, many pregnant women may not have smartphones or may not be comfortable using mobile applications directly. Because of this, a system that depends only on desktop computers or patient smartphone apps may not fit the real situation in government PHCs.",
        "The problem this project addresses, therefore, is the absence of a mobile-first and SMS-based maternal care reminder and follow-up platform designed for government Primary Health Centres in Effurun. Such a system is needed to support health workers with easy patient registration and visit tracking on smartphones or PCs, while also helping pregnant women receive timely reminders through simple SMS communication.",
    ]
    objectives = [
        "To build a mobile-first web platform for government PHC staff to manage maternal care appointments and follow-up.",
        "To provide support for the registration of pregnant women and the recording of antenatal and postnatal visit dates.",
        "To implement a reminder module that prepares and sends SMS notifications before scheduled visits.",
        "To provide missed-visit tracking and follow-up support for women who do not attend their appointments.",
        "To make the system usable on smartphones first, while still allowing access on PCs when available.",
        "To provide a simple reporting dashboard for monitoring due visits, missed visits, and follow-up activities.",
    ]
    conceptual = [
        "Maternal care reminder systems are digital tools used to support pregnant women and new mothers by reminding them of important healthcare visits and follow-up actions. These systems are useful because they help reduce missed appointments and improve continuity of care. In this study, the reminder system is not meant to replace medical care. It is meant to support attendance and follow-up.",
        "A mobile-first system is one that is designed first for small screens such as smartphones before being adjusted for larger screens like laptops and desktop computers. In this project, mobile-first design is important because many PHC workers may depend more on smartphones than standard computers during daily service delivery. This makes mobile usability a major part of the system design rather than an afterthought.",
        "SMS-based health communication involves the use of text messages to share reminders, notices, and follow-up information with patients. It is especially useful in settings where internet access is limited or where many people use basic phones instead of smartphones. In this project, SMS is important because it allows the system to reach women who may not use apps or online services directly.",
        "Follow-up in maternal care means checking on women who are due for care or who have missed scheduled visits. Defaulter tracking means identifying women who did not attend their expected appointments and making them visible for follow-up action. This concept is important in the proposed system because missed visits can weaken continuity of care and reduce timely support for pregnant women.",
        "Government Primary Health Centres are community-level public health facilities that often provide the first level of maternal and child healthcare services. Because they are close to the community, they are suitable settings for a reminder and follow-up platform. In this project, the system is specifically limited to government PHCs in Effurun.",
    ]
    literature_rows = [
        ["Oyinlola et al. (2025), Regional variations in prevalence and factors associated with maternal healthcare services utilisation in Nigeria", "To examine maternal healthcare service use and associated factors across Nigeria.", "National analytical study.", "Gives strong Nigerian evidence that maternal care use is still uneven and affected by multiple factors.", "Does not propose a practical digital reminder and follow-up system for PHCs."],
        ["Osanyin et al. (2022), Effects of a mHealth voice messaging intervention on antenatal care utilisation at primary care level in Lagos, Nigeria", "To examine the effect of voice messaging on antenatal care use.", "Cluster randomised trial.", "Shows that mobile messaging can improve maternal care behaviour in a Nigerian setting.", "Focuses on intervention effect, not on a full clinic management platform."],
        ["Olajubu et al. (2020), Effectiveness of a mobile health intervention on uptake of recommended postnatal care services in Nigeria", "To evaluate the impact of a mobile health intervention on postnatal care use.", "Intervention study.", "Provides Nigerian evidence that mobile health can improve postnatal attendance.", "Does not focus on a mobile-first staff dashboard for PHCs."],
        ["Okonofua et al. (2023), Texting for life: a mobile phone application to connect pregnant women with emergency transport and obstetric care in rural Nigeria", "To support pregnant women through mobile phone linkage to maternal care services.", "Implementation study.", "Shows that mobile phone tools can support maternal care coordination in Nigeria.", "Focuses on emergency linkage rather than routine appointment reminders and follow-up."],
        ["Kante and Malqvist (2025), Effectiveness of SMS-based interventions in enhancing antenatal care in developing countries: a systematic review", "To review the value of SMS interventions for antenatal care.", "Systematic review.", "Gives broad support for SMS reminders in low-resource settings.", "Does not focus specifically on government PHCs in Effurun or staff mobile-first workflow."],
    ]
    methodology = [
        "This project will adopt an incremental software development methodology. This approach is suitable because it allows the system to be developed in stages, tested gradually, and improved as each major part is completed.",
        "The methodology will begin with requirement gathering and study of maternal care workflow in government PHCs. This will involve identifying the main users of the system, the key records needed for maternal appointment tracking, and the reminder and follow-up problems the system is meant to solve.",
        "The next stage will involve system analysis and design. At this point, the system structure, database, interfaces, and reminder workflow will be planned. The proposed system will include a mobile-first dashboard for health workers, patient registration pages, visit scheduling records, a missed-visit tracking section, and an SMS reminder module.",
        "The implementation stage will focus on building the responsive web application and connecting it to a database for storing patient and appointment records. The prototype can be developed using tools such as Next.js, React, TypeScript, a relational database such as SQLite or MySQL, and an SMS gateway or test SMS integration environment for reminder delivery.",
        "The final stage will be testing and evaluation. The system will be tested to confirm that health workers can register women with smartphones, scheduled visits are stored correctly, reminder messages are generated properly, missed visits are flagged clearly, and the dashboard remains easy to use on both phones and PCs.",
    ]
    expected_items = [
        "Mobile-First PHC Dashboard A mobile-first web platform for government PHC staff to register pregnant women, manage maternal visit schedules, and track follow-up actions.",
        "SMS Reminder Module A system component that prepares and sends reminder messages before antenatal and postnatal visits.",
        "Missed-Visit and Follow-Up Tracking A simple dashboard for identifying women due for visits, women who missed appointments, and women who need follow-up.",
        "Improved Maternal Care Support A practical digital tool that can help improve attendance, strengthen follow-up, and support better record management in government PHCs in Effurun.",
    ]
    work_items = [
        "Requirement Gathering and PHC Workflow Study Review of maternal care appointment process, follow-up practices, and reminder needs in government PHCs.",
        "System Analysis and Design Design of the database, mobile-first interface, reminder workflow, and dashboard structure.",
        "Implementation Development of the staff dashboard, patient registration module, appointment tracking features, and SMS reminder functions.",
        "System Testing Evaluation of registration flow, reminder generation, follow-up tracking, and mobile usability.",
        "Documentation and Reporting Preparation of final report, screenshots, findings, and recommendations.",
    ]
    budget_rows = [
        ("Internet/Data", "Data for research, development, testing, and online access", "20,000"),
        ("Power Support", "Electricity and backup power during development", "20,000"),
        ("Transportation", "Visits for project meetings, logistics, and printing", "15,000"),
        ("Printing and Binding", "Draft printing, final printing, and binding", "20,000"),
        ("SMS Testing", "Test message charges and reminder simulation costs", "10,000"),
        ("Hosting/Domain (Optional)", "Optional hosting or domain for demonstration", "15,000"),
        ("Data Backup/Storage", "Backup support for project files and records", "8,000"),
        ("Stationery", "Project stationery and related materials", "5,000"),
        ("User Testing/Logistics", "Small expenses for local testing and coordination", "10,000"),
        ("Contingency", "Unexpected expenses", "12,000"),
        ("Total", "", "135,000"),
    ]
    references = [
        "Hailemariam, T., Atnafu, A., Gezie, L. D., & Tilahun, B. (2024). Effect of short message service reminders in improving optimal antenatal care, skilled birth attendance and postnatal care in low- and middle-income countries: A systematic review and meta-analysis. BMC Medical Informatics and Decision Making, 25, 1. https://doi.org/10.1186/s12911-024-02836-1",
        "Kante, M., & Malqvist, M. (2025). Effectiveness of SMS-based interventions in enhancing antenatal care in developing countries: A systematic review. BMJ Open, 15(2), e089671. https://doi.org/10.1136/bmjopen-2024-089671",
        "Okonofua, F., Ntoimo, L., Johnson, E., Sombie, I., Ojuolape, S., Igboin, B., ... & Yaya, S. (2023). Texting for life: A mobile phone application to connect pregnant women with emergency transport and obstetric care in rural Nigeria. BMC Pregnancy and Childbirth, 23, 139. https://doi.org/10.1186/s12884-023-05424-9",
        "Olajubu, A. O., Fajemilehin, B. R., Olajubu, T. O., & Afolabi, B. S. (2020). Effectiveness of a mobile health intervention on uptake of recommended postnatal care services in Nigeria. PLOS ONE, 15(9), e0238911. https://doi.org/10.1371/journal.pone.0238911",
        "Oluwatola, T., Isiaka, S. D., Omeje, O., Oni, F., Samuel, O. W., Sampson, S., Ebinim, H., & Olatunji, O. (2025). Assessment of quality of maternal and newborn care and its determinants: A national study of primary health care facilities in Nigeria. BMC Health Services Research, 25, 921. https://doi.org/10.1186/s12913-025-12957-6",
        "Osanyin, G. E., Banke-Thomas, A., Oluwole, E. O., Odeseye, A. K., & Afolabi, B. B. (2022). Effects of a mHealth voice messaging intervention on antenatal care utilisation at primary care level in Lagos, Nigeria: A cluster randomised trial. Journal of Public Health in Africa, 13(3), 2222. https://doi.org/10.4081/jphia.2022.2222",
        "Oyinlola, F. F., Okorafor, K. A., Kupoluyi, J. A., Ogbeye, G. B., Ouedraogo, L., Umar, L., & Shittu, I. O. (2025). Regional variations in prevalence and factors associated with maternal healthcare services utilisation in Nigeria. BMC Women's Health, 26, 45. https://doi.org/10.1186/s12905-025-04216-x",
        "Rahman, S., Okolie, A., Bryant, D., Ameyaw, E. K., & Ezezika, O. (2025). Barriers and facilitators of messaging platforms as a means of maternal support and care in rural communities: A systematic review. PLOS ONE, 20(12), e0336168. https://doi.org/10.1371/journal.pone.0336168",
        "UNICEF. (n.d.). Situation of women and children in Nigeria. https://www.unicef.org/nigeria/situation-women-and-children-nigeria",
    ]

    set_para_text(doc.paragraphs[5], "Name: AZETA DUKE")
    set_para_text(doc.paragraphs[6], "Matric Number: COS/8650/2021 Supervisor/Mentor: Dr NWOZOR BLESSING Department: COMPUTER SCIENCE")
    set_para_text(doc.paragraphs[7], "Institution: FEDERAL UNIVERSITY OF PETROLEUM RESOURCES Academic Session: 2025/2026")
    set_para_text(doc.paragraphs[10], title)

    set_para_text(doc.paragraphs[14], executive_summary[0])
    set_para_text(doc.paragraphs[16], executive_summary[1])

    for index, text in zip([19, 21, 23, 25], background):
        set_para_text(doc.paragraphs[index], text)
    clear_para(doc.paragraphs[27])
    clear_para(doc.paragraphs[29])
    background_section = collect_between(doc, "BACKGROUND STUDY", "PROBLEM STATEMENT")
    background_non_empty = [p for p in background_section if p.text.strip()]
    for paragraph in background_non_empty[4:]:
        delete_paragraph(paragraph)

    for index, text in zip([32, 34, 35], problem):
        set_para_text(doc.paragraphs[index], text)

    set_para_text(doc.paragraphs[38], "To design and implement a mobile-first SMS-based maternal care reminder and follow-up platform for government Primary Health Centres in Effurun. The specific objectives of the study are as follows:")

    for index, text in zip([44, 46, 48, 50, 52], conceptual):
        set_para_text(doc.paragraphs[index], text)
    for index in [45, 47, 49, 51, 53, 54]:
        if index < len(doc.paragraphs):
            clear_para(doc.paragraphs[index])

    target_rows = [doc.tables[0].rows[1], doc.tables[0].rows[2], doc.tables[1].rows[0], doc.tables[1].rows[1], doc.tables[1].rows[2]]
    for row, values in zip(target_rows, literature_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    set_para_text(doc.paragraphs[62], "The reviewed studies show that maternal healthcare utilisation, appointment adherence, and mobile health communication are closely related. They also show that messaging tools can improve attendance and continuity of care. However, most of the studies focus either on maternal care use, messaging interventions, or emergency communication. Very few combine these ideas into one mobile-first platform that allows PHC staff to register women, track visits, identify missed appointments, and manage SMS-based follow-up from smartphones or PCs.")
    set_para_text(doc.paragraphs[64], "This project addresses that gap by proposing a mobile-first SMS-based maternal care reminder and follow-up platform for government Primary Health Centres in Effurun. The system combines smartphone-friendly staff registration, antenatal and postnatal scheduling, SMS reminders for women with any type of phone, missed-visit tracking, follow-up support, and a simple dashboard for monitoring due visits and defaulters. This makes the proposed system more practical for the realities of government PHCs.")

    section_paragraphs = collect_between(doc, "AIM/OBJECTIVE", "CONCEPTUAL REVIEW")
    objective_paragraphs = [p for p in section_paragraphs if p.style.name == "List Paragraph"]
    anchor = objective_paragraphs[-1] if objective_paragraphs else doc.paragraphs[38]
    while len(objective_paragraphs) < len(objectives):
        anchor = insert_paragraph_after(anchor, style=doc.paragraphs[39].style)
        objective_paragraphs.append(anchor)
    for paragraph, text in zip(objective_paragraphs[: len(objectives)], objectives):
        set_para_text(paragraph, text)
    for paragraph in objective_paragraphs[len(objectives) :]:
        delete_paragraph(paragraph)

    methodology_heading = find_para(doc, "METHODOLOGY")
    methodology_section = collect_between(doc, "METHODOLOGY", "EXPECTED RESULT / OUTCOME")
    methodology_style = None
    for paragraph in methodology_section:
        if paragraph.text.strip():
            methodology_style = paragraph.style
            break
    if methodology_style is None:
        methodology_style = doc.paragraphs[67].style
    for paragraph in reversed(methodology_section):
        delete_paragraph(paragraph)
    last_methodology = methodology_heading
    for text in methodology:
        last_methodology = insert_paragraph_after(last_methodology, text=text, style=methodology_style)
    methodology_section = collect_between(doc, "METHODOLOGY", "EXPECTED RESULT / OUTCOME")
    methodology_non_empty = [p for p in methodology_section if p.text.strip()]
    for paragraph in methodology_non_empty[len(methodology) :]:
        delete_paragraph(paragraph)

    expected_heading_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "EXPECTED RESULT / OUTCOME")
    set_para_text(doc.paragraphs[expected_heading_index + 1], "The proposed study is expected to deliver the following key outcomes:")
    expected_list_paragraphs = []
    for paragraph in doc.paragraphs[expected_heading_index + 1 :]:
        if paragraph.style.name == "List Paragraph":
            expected_list_paragraphs.append(paragraph)
            if len(expected_list_paragraphs) == 4:
                break
    for paragraph, text in zip(expected_list_paragraphs, expected_items):
        set_para_text(paragraph, text)

    work_heading_index = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "WORK-PLAN/TIME FRAME USING GANTT CHART")
    set_para_text(doc.paragraphs[work_heading_index + 1], "This research will span for a period of 3 months spanning from March 2026 to May 2026.")
    work_table = doc.tables[2]
    work_table.rows[0].cells[0].text = "Description"
    work_table.rows[0].cells[1].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[2].text = "Time Line from March 2026 - May 2026"
    work_table.rows[0].cells[3].text = "Time Line from March 2026 - May 2026"
    work_table.rows[1].cells[0].text = "Description"
    work_table.rows[1].cells[1].text = "MARCH"
    work_table.rows[1].cells[2].text = "APRIL"
    work_table.rows[1].cells[3].text = "MAY"
    work_activities = [("Requirement Gathering/PHC Workflow Study", "X", "", ""), ("System Analysis and Design", "X", "", ""), ("Implementation", "", "X", ""), ("Testing, Evaluation, and Report Writing", "", "", "X")]
    for row_index, values in enumerate(work_activities, start=2):
        for column_index, value in enumerate(values):
            work_table.rows[row_index].cells[column_index].text = value
    set_para_text(doc.paragraphs[work_heading_index + 3], "The implementation of this project will follow a structured workflow consisting of the following stages:")
    work_list_paragraphs = []
    for paragraph in doc.paragraphs[work_heading_index + 4 :]:
        if paragraph.style.name == "List Paragraph":
            work_list_paragraphs.append(paragraph)
            if len(work_list_paragraphs) == 5:
                break
    for paragraph, text in zip(work_list_paragraphs, work_items):
        set_para_text(paragraph, text)
    intro_text = "The implementation of this project will follow a structured workflow consisting of the following stages:"

    for row, values in zip(doc.tables[3].rows[1:], budget_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    references_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"REFERNCES", "REFERENCES"}:
            references_heading = paragraph
            break
    if references_heading is None:
        raise ValueError("References heading not found")
    set_para_text(references_heading, "REFERENCES")
    while True:
        current_paragraphs = doc.paragraphs
        heading_index = next((i for i, paragraph in enumerate(current_paragraphs) if paragraph.text.strip() == "REFERENCES"), -1)
        if heading_index == -1 or heading_index + 1 >= len(current_paragraphs):
            break
        delete_paragraph(current_paragraphs[heading_index + 1])
    last_reference = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "REFERENCES")
    for reference in references:
        last_reference = insert_paragraph_after(last_reference, text=reference, style=references_heading.style)

    intro_hits = [p for p in doc.paragraphs if p.text.strip() == intro_text]
    for paragraph in reversed(intro_hits[1:]):
        delete_paragraph(paragraph)

    banned_phrases = [
        "parallel bilingual corpus",
        "pre-trained T5 transformer",
        "English–Itsekiri",
        "Figure 1: Translation process overview",
        "Figure 2: The Architecture",
        "For the purpose of this work",
        "T5 functions as a sequence-to-sequence model",
        "Upon completion of training and fine-tuning",
        "Through this platform, users will input English text",
    ]
    for paragraph in reversed(list(doc.paragraphs)):
        text = paragraph.text.strip()
        if any(phrase in text for phrase in banned_phrases):
            delete_paragraph(paragraph)

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
