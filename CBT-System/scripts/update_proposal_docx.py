from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TOPIC_DIR = Path(__file__).resolve().parents[1]
ROOT_DIR = TOPIC_DIR.parent
TEMPLATE_PATH = ROOT_DIR / "Shared" / "References" / "Project Proposal.docx"
BACKUP_PATH = ROOT_DIR / "Shared" / "References" / "Project Proposal.original.docx"
DRAFT_PATH = TOPIC_DIR / "Proposal Draft.md"
OUTPUT_PATH = TOPIC_DIR / "Proposal - Azeta Duke.docx"

TITLE = (
    "Design and Implementation of a Mobile-First Anti-Cheat CBT Platform for the "
    "Department of Computer Science, Federal University of Petroleum Resources, Effurun"
)

MAJOR_SECTION_HEADINGS = {
    "EXECUTIVE SUMMARY",
    "BACKGROUND STUDY",
    "PROBLEM STATEMENT",
    "AIM/OBJECTIVES",
    "CONCEPTUAL REVIEW",
    "LITERATURE REVIEW",
    "GAPS IDENTIFIED VIS-A-VIS TECHNOLOGY/METHOD TO RESOLVE THE IDENTIFIED GAPS",
    "METHODOLOGY",
    "EXPECTED RESULT / OUTCOME",
    "WORK-PLAN / TIME FRAME USING GANTT CHART",
    "BUDGET",
    "REFERENCES",
}


def has_drawing(run) -> bool:
    xml = run._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def set_para_text(paragraph, text: str) -> None:
    text_runs = [run for run in paragraph.runs if not has_drawing(run)]
    if text_runs:
        text_runs[0].text = text
        for run in text_runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def remove_body_after_cover(doc: Document, keep_paragraphs: int = 13) -> None:
    body = doc._body._element
    keep_elements = {paragraph._p for paragraph in doc.paragraphs[:keep_paragraphs]}
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        if child not in keep_elements:
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def add_hyperlink(paragraph, text: str, url: str, *, italic: bool = False, bold: bool = False) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    if italic:
        run_properties.append(OxmlElement("w:i"))
    if bold:
        run_properties.append(OxmlElement("w:b"))

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        cursor = match.end()

    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def style_paragraph(
    paragraph,
    *,
    heading_level: int | None = None,
    list_item: bool = False,
    page_break_before: bool = False,
) -> None:
    paragraph.paragraph_format.page_break_before = page_break_before

    if heading_level == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    elif heading_level == 3:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not list_item else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = None

    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if heading_level is not None:
            run.bold = True


def set_document_style(doc: Document) -> None:
    for style_name in ["Normal", "Body Text", "List Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            add_inline_markdown(paragraph, value)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(9 if len(rows[0]) >= 5 else 10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if row_index == 0:
                    run.bold = True

            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")

    doc.add_paragraph()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1

    rows = []
    for table_line in table_lines:
        cells = [cell.strip() for cell in table_line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start_index = next(i for i, line in enumerate(lines) if line.strip() == "## Executive Summary")
    index = start_index

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if line.startswith("|"):
            rows, index = parse_markdown_table(lines, index)
            add_table(doc, rows)
            continue

        if line.startswith("## "):
            text = line[3:].strip().upper()
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, text)
            style_paragraph(
                paragraph,
                heading_level=2,
                page_break_before=text in MAJOR_SECTION_HEADINGS,
            )
            index += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, text)
            style_paragraph(paragraph, heading_level=3)
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        bullet = re.match(r"^-\s+(.*)$", line)
        if numbered or bullet:
            text = numbered.group(1) if numbered else bullet.group(1)
            paragraph = doc.add_paragraph(style="List Paragraph")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            prefix = f"{line.split('.', 1)[0]}. " if numbered else "- "
            paragraph.add_run(prefix)
            add_inline_markdown(paragraph, text)
            style_paragraph(paragraph, list_item=True)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, line)
        style_paragraph(paragraph)
        index += 1


def format_existing_front_matter(doc: Document) -> None:
    for paragraph in doc.paragraphs[:13]:
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)


def main() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)

    markdown = DRAFT_PATH.read_text(encoding="utf-8")
    doc = Document(TEMPLATE_PATH)
    set_document_style(doc)

    set_para_text(doc.paragraphs[5], "Name: AZETA DUKE")
    set_para_text(
        doc.paragraphs[6],
        "Matric Number: COS/8650/2021 Supervisor/Mentor: Dr ABERE Department: COMPUTER SCIENCE",
    )
    set_para_text(
        doc.paragraphs[7],
        "Institution: FEDERAL UNIVERSITY OF PETROLEUM RESOURCES, EFFURUN Academic Session: 2025/2026",
    )
    set_para_text(doc.paragraphs[10], f"Project Title: {TITLE}")

    format_existing_front_matter(doc)
    remove_body_after_cover(doc)
    add_markdown_body(doc, markdown)
    doc.save(OUTPUT_PATH)

    print(f"Created {OUTPUT_PATH}")
    print(f"Source draft: {DRAFT_PATH}")
    print(f"Template backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
