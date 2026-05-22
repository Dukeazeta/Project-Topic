from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TOPIC_DIR = Path(__file__).resolve().parents[1]
INPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 01 Draft.md"
OUTPUT_PATH = TOPIC_DIR / "drafts" / "Chapter 01 - Azeta Duke.docx"

SECTION_HEADING_SPACE_BEFORE = Pt(8)
BODY_AFTER_HEADING_SPACE_BEFORE = Pt(11)


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = None
    normal.paragraph_format.left_indent = Inches(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    pg_num.set(qn("w:start"), "1")
    pg_num.set(qn("w:fmt"), "decimal")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(30) if level == 1 else SECTION_HEADING_SPACE_BEFORE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2.0
    heading_text = text.upper()
    run = paragraph.add_run(heading_text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)


def add_body_paragraph(doc: Document, text: str, *, after_heading: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = BODY_AFTER_HEADING_SPACE_BEFORE if after_heading else Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2.0
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_list_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2.0
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_term_paragraph(doc: Document, line: str) -> None:
    term, meaning = line.split(":", 1)
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2.0
    run = paragraph.add_run(term + ":")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run2 = paragraph.add_run(meaning)
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2.font.size = Pt(12)


def main() -> None:
    doc = Document()
    set_default_font(doc)
    set_margins(doc)
    add_page_number_footer(doc)

    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()

    previous_was_heading = False

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            previous_was_heading = True
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            previous_was_heading = True
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            previous_was_heading = True
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
            previous_was_heading = True
        elif line[:2].isdigit() and line[1] == ".":
            add_list_paragraph(doc, line)
            previous_was_heading = False
        elif line.startswith("**") and ":**" in line:
            cleaned = line.replace("**", "")
            add_term_paragraph(doc, cleaned)
            previous_was_heading = False
        else:
            add_body_paragraph(doc, line, after_heading=previous_was_heading)
            previous_was_heading = False

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
